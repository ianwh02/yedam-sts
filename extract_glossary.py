#!/usr/bin/env python
"""Extract text from Nazarene glossary docx files."""
import io
import sys

import docx

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

for fname in ['nazarene_korean_glossary.docx', 'nazarene_glossary.docx']:
    print(f"\n=== {fname} ===")
    doc = docx.Document(fname)
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text)
    # Also check tables
    for ti, table in enumerate(doc.tables):
        print(f"\n--- Table {ti} ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            print(" | ".join(cells))
